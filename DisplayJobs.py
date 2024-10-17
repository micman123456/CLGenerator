from jobspy import scrape_jobs
import time
import threading
import json
import os
import subprocess
import requests
import docx
from assets.PrompGen import Construct_llama_Prompt,sendAIRequest

# Global Vars
ID = 0
SITE = 1
URL = 2
TITLE = 4
COMPANY = 5
LOCATION = 6
DATE = 8
SALARY_TYPE = 10
SALARY_RANGE_L = 11
SALARY_RANGE_U = 12
CURRENCY = 12
DESC = 20

def validate_input(prompt, valid_options=None, allow_empty=False):
    while True:
        user_input = input(prompt).strip()
        if user_input or allow_empty:
            if valid_options:
                if user_input.lower() in valid_options:
                    return user_input
                else:
                    print(f"Invalid input. Please choose from {valid_options}.")
            else:
                return user_input
        else:
            print("This field cannot be empty. Please enter a value.")

def validate_number(prompt, min_val=None, max_val=None):
    while True:
        try:
            user_input = int(input(prompt).strip())
            if (min_val is not None and user_input < min_val) or (max_val is not None and user_input > max_val):
                print(f"Please enter a number between {min_val} and {max_val}.")
            else:
                return user_input
        except ValueError:
            print("Invalid input. Please enter a valid number.")


def print_dots():
    dots = ""
    
    while not stop_event.is_set():  
        if len(dots) < 3:
            dots += "."
        else:
            dots = ""
            print(f"\r   ", end='', flush=True) 
            
        print(f"\r{dots}", end='', flush=True)  
        time.sleep(0.5) 

def start_ollama():
    print("Starting ollama, please wait...\n")
    
    process = subprocess.Popen(["ollama", "serve"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    
    time.sleep(2)

    
    for _ in range(10):
        try:
            response = requests.get('http://localhost:11434/')
            if response.text == 'Ollama is running':
                print("Ollama is running")
                return True
        except requests.ConnectionError:
            pass
        
        time.sleep(2) 

   
    print("Please install and run Ollama to continue")
    process.terminate()  
    return False

def collect_user_data():
    data = {}
    data['name'] = input("Enter your Full Name: \n")
    data['email'] = input("Enter your Email:\n")
    data['number'] = input("Enter your Phone Number: \n")
    data['location'] = input("Enter location. City, State, Country:\n")
    data['skills'] = input("Enter a list of skills separated by commas:\n").split(",")
    data['languages'] = input("Enter a list of programming languages separated by commas:\n").split(",")
    data['experience'] = input("Enter a list of previous jobs separated by commas:\n").split(",")

    print("Enter the number corresponding to your highest level of education")
    print("1. High School 2. Community College 3. Undergraduate Degree 4. Master's Degree")
    uinput = input("Enter the number corresponding to your education level: \n")
    if not uinput.isdigit():
        uinput = 0
    else:
         uinput = int(uinput)

    match uinput:
        case 1:
            data['education'] = "High School Diploma"
        case 2:
            data['education'] = "Community College Diploma/Certificate"
        case 3:
            data['education'] = "Undergraduate Degree"
        case 4:
            data['education'] = "Master's Degree"
        case _:
            data['education'] = "Idiot's Degree"

    if uinput > 1 and uinput < 5:
        data['field'] = input("What was your field of study: \n")
    else:
        data['field'] = "General Education"

    data['school'] = input("What is the name of the establishment you attended?\n")
    data['certs'] = input("If you have any other certifications, enter them, each separated by commas: \n").split(",")

    return data

def save_user_data(data):
    with open('user/user.json', 'w') as file:
        json.dump(data, file, indent=4)
    print("Data saved successfully.")


def write_jobs_to_html(jobs, file_name, title="Job Listings"):
    with open(file_name, 'w', encoding='utf-8') as file:
        # Start the HTML structure
        file.write("<html>\n<head>\n<title>Job Listings</title>\n</head>\n<body>\n")
        file.write(f"<h1>{title}</h1>\n")
        
        # Begin the unordered list
        file.write("<ul>\n")
        
        for _, job in jobs.iterrows():
            # Write each job with details in <li> elements using iloc for positional access
            file.write(f"  <li>\n")
            file.write(f"<p><strong>Job ID:</strong> {job.iloc[ID]}</p>\n")
            file.write(f"<p><strong>Job Site:</strong> {job.iloc[SITE]}</p>\n")
            file.write(f"<p><strong>Title:</strong> {job.iloc[TITLE]}</p>\n")
            file.write(f"<p><strong>Company:</strong> {job.iloc[COMPANY]}</p>\n")
            file.write(f"<p><strong>Location:</strong> {job.iloc[LOCATION]}</p>\n")
            file.write(f"<p><strong>Posted Date:</strong> {job.iloc[DATE]}</p>\n")
            file.write(f"<p><strong>Salary Type:</strong> {job.iloc[SALARY_TYPE]}</p>\n")
            file.write(f"<p><strong>Salary Range:</strong> {job.iloc[SALARY_RANGE_L]} - {job.iloc[SALARY_RANGE_U]} {job.iloc[CURRENCY]}</p>\n")
            
            # Write job description
            file.write(f"<p><strong>Description:</strong></p>\n")
            job_desc = job.iloc[DESC]

            if isinstance(job_desc, str):  
                job_desc_list = job_desc.split("**")  
                for row in job_desc_list:
                    file.write(f"<p>{row}</p>\n") 
            else:
                file.write("<p>No job description available</p>\n")

            file.write(f"<p><a href='{job.iloc[URL]}'>Job URL</a></p>\n")
            file.write(f"  </li>\n")  # End the <li> for each job
            
            file.write(f"<br>\n")  # Add a line break after each job entry

        # End the unordered list and close the HTML structure
        file.write("</ul>\n</body>\n</html>")


def print_jobs_to_console(jobs):
    index = 0
    for _, job in jobs.iterrows():
        print(f'{index}: {job.iloc[TITLE]} at {job.iloc[COMPANY]} in {job.iloc[LOCATION]} posted on {job.iloc[DATE]}')
        print(f'link : {job.iloc[URL]}\n')
        index += 1



def main():

    # Call the function to start ollama
    if not start_ollama():
        exit(1)


    print("*** Welcome to Job Finder ***\n")




    
    with open('user/user.json', 'r') as file:
    
        if os.stat('user/user.json').st_size == 0:
            data = {}
            print("No user data found. Please enter the following information to continue\n")
        else:
            data = json.load(file)

    
    if not data:
        user_data = collect_user_data()
        save_user_data(user_data)
        
        # Attempt to fetch again.
        with open('user/user.json', 'r') as file:
    
            if os.stat('user/user.json').st_size == 0:
                data = {}
                print("An Error has occurred. Please recreate the user.json file.\n")
                return
            else:
                data = json.load(file)
        print("Data saved successfully.")

        
            
        

    print("Please enter a job title or keyword")
    jobTitle = validate_input("Job title/keyword: ")

    print("Please enter your Country")
    country = validate_input("Country: ")

    print("Please enter your Province/State (in abbreviated form). E.g., Nova Scotia = NS")
    state = validate_input("Province/State: ")

    print("Please enter your City")
    city = validate_input("City: ")

    # Uncomment if you want to try remote jobs but the job scraper kinda tweaks out when I include remote jobs so

    # print("Do you want to search for remote jobs as well? (yes/no)")
    # include_remote = validate_input("")
    # if (include_remote.lower() == "yes"):
    #     include_remote = True
    # else:
    #     include_remote = False
    include_remote = False

    print("Please enter the sites you want to search (separated by space). Enter 'all' for all sites (indeed linkedin zip_recruiter glassdoor)")
    sites_input = validate_input("Sites: ", allow_empty=False)
    if sites_input.lower() != "all":
        sites = sites_input.split()
    else:
        sites = ["indeed", "linkedin", "zip_recruiter", "glassdoor"]

    results_wanted = validate_number("Please enter the number of results per site you want (max 20): ", min_val=1, max_val=20)

    print("Searching " + " ".join(str(x) for x in sites) + " with the following parameters\n")
    print(f"Search term : {jobTitle} \nLocation : {city}, {state}, {country} {'(Remote included)' if include_remote else ''}\n")

    # Call to job scraper
    try:
        jobs = scrape_jobs(
            site_name=sites,
            search_term=jobTitle,
            location=city + ", " + state,
            is_remote=include_remote,
            results_wanted=results_wanted,
            hours_old=200,  # (only LinkedIn/Indeed is hour-specific, others round up to days old)
            country_indeed=country,  # only needed for indeed/glassdoor
              
        )
    except Exception as e:
        print(f"Error during job search: {e}")

    # Print how many jobs were found
    print(f"Found {len(jobs)} jobs\n")
    
    print_jobs_to_console(jobs)


    # Write jobs to an HTML file with UTF-8 encoding
    write_jobs_to_html(jobs, 'html/jobs.html')
        
    print("Jobs successfully written to 'jobs.html'")

    regenerate = False

    while True:
        
        if not regenerate:

            print("Enter the job # below to generate a Cover Letter for that job. Otherwise type 'exit' to quit")
            uinput = input()
            if (uinput == "exit"):
                return
            elif (not uinput.isdigit()):
                return
                
            
            jobNumber = int(uinput)
            
            job = jobs.iloc[jobNumber]

        print(f"Generating cover letter for {job.iloc[TITLE]} at {job.iloc[COMPANY]}")
        print(f"Please wait...")

        prompt = Construct_llama_Prompt(data, jobs.iloc[jobNumber])
        
        dot_thread = threading.Thread(target=print_dots)
        dot_thread.start()

        
        response = sendAIRequest(prompt)
        
        
        stop_event.set()
        dot_thread.join() 

        print("Would you like to save that cover letter? yes/no")
        uinput = str(input())
        if (uinput == 'yes'):
            doc = docx.Document() 
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_after = 0

            # Add a run to the paragraph
            p.add_run(response)
            doc.save(f"cover_letters/CoverLetter_{job.iloc[COMPANY]}.docx")
            regenerate = False
            
        
        else:
            print("Would you like to regenerate the cover letter? yes/no")
            uinput = str(input())
            if (uinput == 'yes'):
                regenerate = True
            else:
                regenerate = False


            
        
    
stop_event = threading.Event()

if __name__=="__main__":
    main()